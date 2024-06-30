from docx import Document
import re

# 用来提取实验目的和实验记录
def extract_specific_sections_from_docx(docx_file):
    doc = Document(docx_file)
    sections_dict = {"试验目的": "", "试验记录": ""}
    current_section = None

    skip_section = False  # 标记是否跳过当前一级标题或二级标题下的正文内容

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # 提取标题核心内容，忽略任何数字或其他字符
        title_match = re.match(r'^\s*\d*\s*(.+)', text)  # 匹配开头可能存在的数字和空格
        if title_match:
            text = title_match.group(1).strip()

        # 检查是否是我们需要的一级标题
        if text in sections_dict:
            current_section = text
            skip_section = False  # 进入目标段落，不跳过内容
        elif current_section and (paragraph.style.name.startswith("Heading 1") or paragraph.style.name.startswith("Heading 2")):
            skip_section = True  # 一级或二级标题，标记跳过内容
        elif current_section and not skip_section:
            # 将内容追加到当前部分，并保留换行符
            # sections_dict[current_section] += paragraph.text + '\n'
            sections_dict[current_section] += paragraph.text

    # 移除每部分内容末尾多余的换行符和标题本身
    # sections_dict = {k: v.split(k)[-1].strip() for k, v in sections_dict.items()}
    sections_dict = {k: v.strip() for k, v in sections_dict.items()}

    return sections_dict

# 用来提取表格中的非结构化数据
def read_tables_from_docx(docx_file):
    doc = Document(docx_file)
    all_tables = {}

    for table_index, table in enumerate(doc.tables):
        table_name = f'数据_{table_index + 1}'  # 自定义表格名称
        table_dict = {}

        # 遍历表格的行，跳过第一行（假设为表头）
        for row_index, row in enumerate(table.rows[1:], start=1):
            if len(row.cells) >= 2:
                key = row.cells[1].text.strip()  # 假设键位于第二列
                values = [cell.text.strip() for cell in row.cells[2:]]  # 第三列及之后列的值
                table_dict[key] = values

        all_tables[table_name] = table_dict

    return all_tables

# 示例用法
docx_file = 'D:/test2.docx'  # 替换为你的 Word 文件路径

try:
    sections_dict = extract_specific_sections_from_docx(docx_file)
    tables_dict = read_tables_from_docx(docx_file)

    print(sections_dict)
    print(tables_dict)

    # for table_name, table_data in tables_dict.items():
    #     print(f"\n表格: {table_name}")
    #     for key, values in table_data.items():
    #         print(f"{key}：{values}")


except Exception as e:
    print(f"出现错误：{e}")
